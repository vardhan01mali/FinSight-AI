import docx
import logging

logger = logging.getLogger(__name__)

def parse_docx(file_path: str) -> list:
    """
    Parses a DOCX document using python-docx.
    Simulates page numbers by grouping text blocks into roughly 3000 character pages.
    """
    try:
        doc = docx.Document(file_path)
        full_text_blocks = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text_blocks.append(para.text.strip())
                
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    full_text_blocks.append(" | ".join(row_cells))
                    
        # Group text into simulated pages (approx 3000 chars per page)
        pages_data = []
        current_page_text = []
        current_char_count = 0
        page_counter = 1
        
        for block in full_text_blocks:
            current_page_text.append(block)
            current_char_count += len(block)
            
            if current_char_count >= 3000:
                pages_data.append({
                    "page": page_counter,
                    "text": "\n\n".join(current_page_text)
                })
                current_page_text = []
                current_char_count = 0
                page_counter += 1
                
        # Append remaining text
        if current_page_text:
            pages_data.append({
                "page": page_counter,
                "text": "\n\n".join(current_page_text)
            })
            
        if not pages_data:
            return [{"page": 1, "text": "Empty Word Document."}]
            
        return pages_data
    except Exception as e:
        logger.error(f"Error parsing DOCX file {file_path}: {e}")
        raise RuntimeError(f"DOCX parsing failed: {str(e)}")


def parse_txt(file_path: str) -> list:
    """
    Parses a plain text file.
    Simulates pages by splitting text every 3000 characters.
    """
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
            
        if not text.strip():
            return [{"page": 1, "text": "Empty text file."}]
            
        chunk_size = 3000
        pages_data = []
        page_counter = 1
        
        for i in range(0, len(text), chunk_size):
            chunk = text[i:i+chunk_size]
            pages_data.append({
                "page": page_counter,
                "text": chunk.strip()
            })
            page_counter += 1
            
        return pages_data
    except Exception as e:
        logger.error(f"Error parsing TXT file {file_path}: {e}")
        raise RuntimeError(f"TXT parsing failed: {str(e)}")
